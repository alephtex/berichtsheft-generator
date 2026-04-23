#!/usr/bin/env python3
"""
Berichtsheft DOCX Updater
Ersetzt Platzhalter in der aktuellen KW-DOCX und passt automatisch die Größe an.

Usage:
    python3 fill_docx.py --kw 16 --jahr 2026 --nr 136 --abteilung Berufsschule \
        --montag "Text" --dienstag "Text" --mittwoch "Text" \
        --donnerstag "Text" --freitag "Text" --thema "Text"

Platzhalter in Vorlage:
    {Nr}, {Ausbildungsjahr}, {KW}, {datums-range}
    {Montag}, {Dienstag}, {Mittwoch}, {Donnerstag}, {Freitag}
    {week_topic}, {Abteilung}, {day_of_signature}
"""

import argparse
import re
import sys
import zipfile
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt
from lxml import etree
from datetime import datetime, timedelta

def replace_in_paragraph(paragraph, replacements):
    """Ersetzt Platzhalter in allen Runs eines Paragraphs (treats runs as one text)."""
    if not paragraph.runs:
        return
    
    # Join all runs
    full_text = ''.join(run.text for run in paragraph.runs)
    
    # Check if any placeholder exists
    has_match = False
    for placeholder in replacements.keys():
        if placeholder in full_text:
            has_match = True
            break
    
    if not has_match:
        return
    
    # Replace
    for placeholder, value in replacements.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, value)
    
    # Clear all runs and set first run
    for run in paragraph.runs:
        run.text = ''
    paragraph.runs[0].text = full_text

def replace_in_cell(cell, replacements):
    """Ersetzt Platzhalter in allen Runs einer Tabellenzelle."""
    for para in cell.paragraphs:
        replace_in_paragraph(para, replacements)

def set_abteilung_font(doc, abteilung):
    """Setzt Font-Size für Abteilung-Zellen auf 8."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip() == abteilung:
                        for run in para.runs:
                            run.font.size = Pt(8)

def get_kw_dates(kw, year):
    """Berechnet Montag und Sonntag einer Kalenderwoche."""
    jan4 = datetime(year, 1, 4)
    iso = jan4.isocalendar()
    days_since_monday = iso[2] - 1
    monday_kw1 = jan4 - timedelta(days=days_since_monday)
    monday = monday_kw1 + timedelta(weeks=kw - 1)
    sunday = monday + timedelta(days=6)
    saturday = monday + timedelta(days=5)
    return monday, sunday, saturday

def has_placeholders(docx_path):
    """Prüft ob noch Platzhalter im Dokument vorhanden sind."""
    with zipfile.ZipFile(docx_path, 'r') as z:
        xml_content = z.read('word/document.xml')
    root = etree.fromstring(xml_content)
    for t in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        if t.text and '{' in t.text:
            return True
    return False

def shrink_xml(docx_path, factor=0.95):
    """Verkleinert Zeilenhöhen im XML."""
    with zipfile.ZipFile(docx_path, 'r') as z:
        xml_content = z.read('word/document.xml')
    
    root = etree.fromstring(xml_content)
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    modified = 0
    for trHeight in root.iter(f'{{{ns}}}trHeight'):
        val = trHeight.get(f'{{{ns}}}val')
        if val and val.isdigit():
            val_int = int(val)
            if val_int > 100:
                new_val = max(100, int(val_int * factor))
                trHeight.set(f'{{{ns}}}val', str(new_val))
                modified += 1
    
    print(f"  Shrunk {modified} Zeilen")
    
    new_xml = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')
    
    tmp_path = docx_path.with_suffix('.tmp.docx')
    with zipfile.ZipFile(docx_path, 'r') as z_in:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as z_out:
            for item in z_in.namelist():
                if item == 'word/document.xml':
                    z_out.writestr(item, new_xml)
                else:
                    z_out.writestr(item, z_in.read(item))
    
    import os
    os.replace(tmp_path, docx_path)

def update_checkboxes(docx_path, abteilung):
    """Setzt Checkboxen direkt im XML."""
    abt_lower = abteilung.lower()
    
    with zipfile.ZipFile(docx_path, 'r') as z:
        xml_content = z.read('word/document.xml')
    
    root = etree.fromstring(xml_content)
    
    # Checkboxen zurücksetzen
    for t in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        if t.text in ['☐', '☒', '[X]', '[ ]']:
            t.text = '☐'
    
    # Finde Checkboxen und deren Position
    checkboxes = []
    for t in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        if t.text == '☐':
            checkboxes.append(t)
    
    # Richtige Checkbox aktivieren
    if abt_lower == 'betrieb' and len(checkboxes) > 0:
        checkboxes[0].text = '☒'
    elif abt_lower == 'berufsschule' and len(checkboxes) > 1:
        checkboxes[1].text = '☒'
    
    new_xml = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')
    
    tmp_path = docx_path.with_suffix('.tmp.docx')
    with zipfile.ZipFile(docx_path, 'r') as z_in:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as z_out:
            for item in z_in.namelist():
                if item == 'word/document.xml':
                    z_out.writestr(item, new_xml)
                else:
                    z_out.writestr(item, z_in.read(item))
    
    import os
    os.replace(tmp_path, docx_path)

def convert_to_pdf(docx_path, pdf_path):
    """Konvertiert DOCX zu PDF mit LibreOffice."""
    import os
    os.chdir(docx_path.parent)
    
    result = subprocess.run(
        ['libreoffice', '--headless', '--convert-to', 'pdf', str(docx_path), '--outdir', str(docx_path.parent)],
        capture_output=True, text=True, timeout=60
    )
    
    expected_pdf = docx_path.with_suffix('.pdf')
    if expected_pdf.exists():
        return True
    
    for f in docx_path.parent.glob('*.pdf'):
        if f.stat().st_mtime > docx_path.stat().st_mtime:
            if f != pdf_path:
                f.rename(pdf_path)
            return True
    
    return False

def check_pages(pdf_path):
    """Gibt Seitenanzahl zurück."""
    result = subprocess.run(
        ['pdfinfo', str(pdf_path)],
        capture_output=True, text=True
    )
    for line in result.stdout.split('\n'):
        if line.startswith('Pages:'):
            return int(line.split(':')[1].strip())
    return -1

def main():
    parser = argparse.ArgumentParser(description='Berichtsheft DOCX Updater')
    parser.add_argument('--kw', type=int, required=True)
    parser.add_argument('--jahr', type=int, required=True)
    parser.add_argument('--nr', type=int, required=True)
    parser.add_argument('--abteilung', type=str, required=True, choices=['Betrieb', 'Berufsschule'])
    parser.add_argument('--montag', type=str, default='')
    parser.add_argument('--dienstag', type=str, default='')
    parser.add_argument('--mittwoch', type=str, default='')
    parser.add_argument('--donnerstag', type=str, default='')
    parser.add_argument('--freitag', type=str, default='')
    parser.add_argument('--thema', type=str, default='')
    parser.add_argument('--ausbildungsjahr', type=int, default=3)
    
    args = parser.parse_args()
    
    # Pfade
    kw_dir = Path.home() / 'Nextcloud' / 'Dokumente' / 'Ausbildung-Cancom' / 'Berichtsheft' / str(args.jahr) / f'KW{args.kw:02d}'
    kw_dir.mkdir(parents=True, exist_ok=True)
    docx_path = kw_dir / f'Berichtsheft Täglich KW {args.kw:02d}.docx'
    pdf_path = kw_dir / f'Berichtsheft Täglich KW {args.kw:02d}.pdf'
    
    if not docx_path.exists():
        print(f"FEHLER: DOCX nicht gefunden: {docx_path}")
        sys.exit(1)
    
    # KW-Datumsbereich berechnen
    monday, sunday, saturday = get_kw_dates(args.kw, args.jahr)
    kw_date_range = f"{monday.day:02d}.{monday.month:02d}.{args.jahr} - {sunday.day:02d}.{sunday.month:02d}.{args.jahr}"
    day_month = f"{saturday.day:02d}.{saturday.month:02d}"
    signature_date = f"{day_month}.{args.jahr}"
    
    # Replacements (Platzhalter aus der Vorlage)
    replacements = {
        '{Nr}': str(args.nr),
        '{Ausbildungsjahr}': str(args.ausbildungsjahr),
        '{KW}': f"{args.kw:02d}",
        '{datums-range}': kw_date_range,
        '{Montag}': args.montag,
        '{Dienstag}': args.dienstag,
        '{Mittwoch}': args.mittwoch,
        '{Donnerstag}': args.donnerstag,
        '{Freitag}': args.freitag,
        '{week_topic}': args.thema,
        '{day_of_signature}': signature_date,
        '{Abteilung}': args.abteilung,
    }
    
    print(f"Öffne: {docx_path}")
    
    doc = Document(str(docx_path))
    
    # In Tables ersetzen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, replacements)
    
    # In Paragraphs ersetzen
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    
    doc.save(str(docx_path))
    
    # Abteilung Font auf 8 setzen
    doc2 = Document(str(docx_path))
    set_abteilung_font(doc2, args.abteilung)
    doc2.save(str(docx_path))
    print(f"✓ Platzhalter ersetzt, Abteilung Font 8")
    
    # Checkboxen aktualisieren
    update_checkboxes(docx_path, args.abteilung)
    print(f"✓ Checkbox '{args.abteilung}' gesetzt")
    
    # PDF konvertieren
    if not convert_to_pdf(docx_path, pdf_path):
        print("⚠️  PDF-Konvertierung fehlgeschlagen")
        sys.exit(0)
    
    pages = check_pages(pdf_path)
    
    if pages == 1:
        print(f"✓ Perfekt: 1 Seite")
        subprocess.run(['xdg-open', str(pdf_path)])
        return
    
    if pages > 1:
        print(f"\n⚠️  {pages} Seiten - Verkleinere...")
        
        factor = 0.95
        shrink_count = 0
        max_shrinks = 10
        
        while shrink_count < max_shrinks:
            shrink_count += 1
            shrink_xml(docx_path, factor=factor)
            update_checkboxes(docx_path, args.abteilung)
            
            if convert_to_pdf(docx_path, pdf_path):
                pages2 = check_pages(pdf_path)
                
                if pages2 == 1:
                    print(f"✓ Perfekt: 1 Seite (nach {shrink_count}. Anpassung)")
                    subprocess.run(['xdg-open', str(pdf_path)])
                    return
                
                print(f"  Immernoch {pages2} Seiten...")
                factor = max(0.80, factor - 0.03)
            else:
                break
        
        print(f"⚠️  PDF hat {pages} Seiten - bitte Text kürzen!")
        subprocess.run(['xdg-open', str(pdf_path)])

if __name__ == '__main__':
    main()
